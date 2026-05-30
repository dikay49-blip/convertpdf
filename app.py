import os, uuid, threading, time, logging, mimetypes
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, send_file, jsonify, abort, send_from_directory, Response
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf.csrf import CSRFProtect
from werkzeug.utils import secure_filename
import filetype, img2pdf
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
import docx
from xhtml2pdf import pisa

app = Flask(__name__, static_folder='static', static_url_path='/static')
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-prod')
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024
app.config['WTF_CSRF_TIME_LIMIT'] = 3600

UPLOAD_FOLDER = Path('uploads')
UPLOAD_FOLDER.mkdir(exist_ok=True)
SITE_URL = os.environ.get('SITE_URL', 'https://convertpdf-ezqw.onrender.com')

logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s')
logger = logging.getLogger(__name__)

csrf = CSRFProtect(app)
limiter = Limiter(get_remote_address, app=app,
default_limits=["1000 per day", "200 per hour"], storage_uri="memory://")

ALLOWED_MIME = {
    'image/jpeg':'image','image/png':'image','image/gif':'image',
    'image/webp':'image','image/bmp':'image','image/tiff':'image',
    'text/plain':'text','text/html':'html',
    'application/msword':'docx',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document':'docx',
    'text/csv':'csv','application/csv':'csv',
}
ALLOWED_EXT = {'.jpg','.jpeg','.png','.gif','.webp','.bmp','.tiff',
               '.txt','.html','.htm','.doc','.docx','.csv'}
TEXT_EXT = {'.txt':'text/plain','.html':'text/html','.htm':'text/html','.csv':'text/csv'}

@app.after_request
def add_security_headers(r):
    r.headers['X-Content-Type-Options']    = 'nosniff'
    r.headers['X-Frame-Options']           = 'SAMEORIGIN'
    r.headers['X-XSS-Protection']          = '1; mode=block'
    r.headers['Referrer-Policy']           = 'strict-origin-when-cross-origin'
    r.headers['Permissions-Policy']        = 'geolocation=(), microphone=(), camera=()'
    r.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    r.headers['Content-Security-Policy']   = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://pagead2.googlesyndication.com "
        "https://www.paypal.com https://www.paypalobjects.com https://cdn.jsdelivr.net "
        "https://partner.googleadservices.com https://tpc.googlesyndication.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://cdn.jsdelivr.net; "
        "font-src 'self' https://fonts.gstatic.com https://cdn.jsdelivr.net; "
        "img-src 'self' data: https: blob:; "
        "frame-src https://www.paypal.com https://pagead2.googlesyndication.com "
        "https://tpc.googlesyndication.com; "
        "connect-src 'self'; object-src 'none'; base-uri 'self';"
    )
    if request.path.startswith('/download/'):
        r.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, private'
    return r

def auto_delete():
    while True:
        try:
            now = time.time()
            for f in UPLOAD_FOLDER.iterdir():
                if f.is_file() and now - f.stat().st_mtime > 300:
                    f.unlink(missing_ok=True)
        except: pass
        time.sleep(60)

threading.Thread(target=auto_delete, daemon=True).start()

def validate_file(fs):
    name = secure_filename(fs.filename or '')
    if not name: return None, None, "Nom invalide."
    ext = Path(name).suffix.lower()
    if ext not in ALLOWED_EXT: return None, None, f"Extension non supportée : {ext}"
    header = fs.read(2048); fs.seek(0)
    kind = filetype.guess(header)
    mime = kind.mime if kind else TEXT_EXT.get(ext)
    if not mime: return None, None, "Type non reconnu."
    if mime not in ALLOWED_MIME: return None, None, f"Type non supporté : {mime}"
    return name, mime, None

def safe_path(folder, filename):
    t = (folder / filename).resolve()
    if not str(t).startswith(str(folder.resolve())): abort(400)
    return t

def img_to_pdf(src, dst):
    with Image.open(src) as img:
        if img.mode in ('RGBA','P','LA'):
            tmp = src.with_suffix('.tmp.jpg')
            img.convert('RGB').save(tmp,'JPEG',quality=95)
            dst.write_bytes(img2pdf.convert(open(tmp,'rb')))
            tmp.unlink(missing_ok=True)
        else:
            dst.write_bytes(img2pdf.convert(open(src,'rb')))

def txt_to_pdf(src, dst):
    doc = SimpleDocTemplate(str(dst), pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    st = ParagraphStyle('b', parent=getSampleStyleSheet()['Normal'],
        fontName='Helvetica', fontSize=11, leading=16)
    items = []
    for line in src.read_text(encoding='utf-8', errors='replace').splitlines():
        line = line.strip()
        if line:
            items.append(Paragraph(line.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'), st))
            items.append(Spacer(1,4))
        else:
            items.append(Spacer(1,12))
    if not items: items.append(Paragraph("(vide)", st))
    doc.build(items)

def html_to_pdf(src, dst):
    with open(dst,'wb') as f:
        pisa.CreatePDF(src.read_text(encoding='utf-8',errors='replace'), dest=f)

def docx_to_pdf(src, dst):
    d = docx.Document(str(src))
    h = ["<html><head><meta charset='utf-8'><style>"
         "body{font-family:Arial;font-size:12pt;margin:2cm;line-height:1.5;}"
         "h1{font-size:18pt;}h2{font-size:15pt;}h3{font-size:13pt;}"
         "table{border-collapse:collapse;width:100%;}td,th{border:1px solid #ccc;padding:5px;}"
         "</style></head><body>"]
    for p in d.paragraphs:
        t = p.text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        if not t.strip(): h.append('<br>'); continue
        if p.style.name.startswith('Heading 1'): h.append(f'<h1>{t}</h1>')
        elif p.style.name.startswith('Heading 2'): h.append(f'<h2>{t}</h2>')
        elif p.style.name.startswith('Heading 3'): h.append(f'<h3>{t}</h3>')
        else: h.append(f'<p>{t}</p>')
    for tbl in d.tables:
        h.append('<table>')
        for i,row in enumerate(tbl.rows):
            h.append('<tr>')
            tag = 'th' if i==0 else 'td'
            for cell in row.cells:
                ct = cell.text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                h.append(f'<{tag}>{ct}</{tag}>')
            h.append('</tr>')
        h.append('</table>')
    h.append('</body></html>')
    with open(dst,'wb') as f:
        pisa.CreatePDF('\n'.join(h), dest=f)

def csv_to_pdf(src, dst):
    import csv
    doc = SimpleDocTemplate(str(dst), pagesize=A4,
        leftMargin=1*cm, rightMargin=1*cm, topMargin=2*cm, bottomMargin=2*cm)
    rows = list(csv.reader(open(src, encoding='utf-8', errors='replace')))
    if not rows:
        doc.build([Paragraph("Vide", getSampleStyleSheet()['Normal'])]); return
    t = Table(rows, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#16213e')),
        ('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
        ('FONTSIZE',(0,0),(-1,-1),9),
        ('GRID',(0,0),(-1,-1),0.5,colors.lightgrey),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#f5f7fa')]),
    ]))
    doc.build([t])

CONV = {'image':img_to_pdf,'text':txt_to_pdf,'html':html_to_pdf,'docx':docx_to_pdf,'csv':csv_to_pdf}

@app.context_processor
def inject_now(): return {'now': datetime.utcnow()}

@app.route('/')
def index(): return render_template('index.html')

@app.route('/donate')
def donate(): return render_template('donate.html')

@app.route('/privacy')
def privacy(): return render_template('privacy.html')

@app.route('/guide')
def guide(): return render_template('guide.html')

@app.route('/about')
def about(): return render_template('about.html')

@app.route('/robots.txt')
def robots():
    return Response(f"User-agent: *\nAllow: /\nDisallow: /convert\nDisallow: /download/\n\nSitemap: {SITE_URL}/sitemap.xml\n", mimetype='text/plain')

@app.route('/sitemap.xml')
def sitemap():
    now = datetime.utcnow().strftime('%Y-%m-%d')
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <url><loc>{SITE_URL}/</loc><lastmod>{now}</lastmod><priority>1.0</priority></url>
  <url><loc>{SITE_URL}/guide</loc><lastmod>{now}</lastmod><priority>0.9</priority></url>
  <url><loc>{SITE_URL}/about</loc><lastmod>{now}</lastmod><priority>0.7</priority></url>
  <url><loc>{SITE_URL}/donate</loc><lastmod>{now}</lastmod><priority>0.5</priority></url>
  <url><loc>{SITE_URL}/privacy</loc><lastmod>{now}</lastmod><priority>0.4</priority></url>
</urlset>"""
    return Response(xml, mimetype='application/xml')

@app.route('/convert', methods=['POST'])
@limiter.limit("10 per minute")
def convert():
    if 'file' not in request.files: return jsonify({'error':'Aucun fichier.'}), 400
    f = request.files['file']
    if not f or not f.filename: return jsonify({'error':'Fichier vide.'}), 400
    name, mime, err = validate_file(f)
    if err: return jsonify({'error':err}), 400
    uid = uuid.uuid4().hex
    ext = Path(name).suffix.lower()
    inp = safe_path(UPLOAD_FOLDER, f'{uid}_in{ext}')
    out = safe_path(UPLOAD_FOLDER, f'{uid}_out.pdf')
    try:
        f.save(str(inp))
        CONV[ALLOWED_MIME[mime]](inp, out)
        inp.unlink(missing_ok=True)
        return jsonify({'download_id':uid, 'original_name':Path(name).stem})
    except Exception as e:
        logger.error(f'Conversion error: {e}', exc_info=True)
        inp.unlink(missing_ok=True); out.unlink(missing_ok=True)
        return jsonify({'error':str(e)}), 500

@app.route('/download/<uid>')
@limiter.limit("20 per minute")
def download(uid):
    if not uid.isalnum() or len(uid) != 32: abort(400)
    out = safe_path(UPLOAD_FOLDER, f'{uid}_out.pdf')
    if not out.exists(): abort(404)
    name = secure_filename(request.args.get('name','document')) or 'document'
    def del_later(p):
        time.sleep(5); Path(p).unlink(missing_ok=True)
    threading.Thread(target=del_later, args=(str(out),), daemon=True).start()
    return send_file(str(out), as_attachment=True, download_name=f'{name}.pdf', mimetype='application/pdf')

@app.errorhandler(400)
def e400(e): return render_template('error.html', code=400, msg="Requête invalide."), 400
@app.errorhandler(404)
def e404(e): return render_template('error.html', code=404, msg="Page introuvable."), 404
@app.errorhandler(413)
def e413(e): return jsonify({'error':'Fichier trop volumineux (max 20 Mo).'}), 413
@app.errorhandler(429)
def e429(e): return jsonify({'error':'Trop de requêtes. Attends 1 minute.'}), 429
@app.errorhandler(500)
def e500(e): return render_template('error.html', code=500, msg="Erreur interne."), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT',5000)), debug=False)
